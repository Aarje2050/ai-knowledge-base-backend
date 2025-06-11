import os
import uuid
import aiofiles
import re
import pdfplumber  # MIT License - Free alternative to PyMuPDF
from typing import List, Dict, Any, Optional
from docx import Document
import nltk
import pytesseract
from PIL import Image
import io

from app.config import settings
from app.services.ai_service import AdvancedAIService
from app.services.vector_store import AdvancedVectorStore

class AdvancedDocumentProcessor:
    def __init__(self):
        self.ai_service = AdvancedAIService()
        self.vector_store = AdvancedVectorStore()
        
        # Chunking parameters
        self.target_chunk_size = 800  # Reduced for better precision
        self.chunk_overlap = 100
        self.max_chunk_size = 1200
        
        # Initialize NLTK for sentence splitting
        try:
            nltk.data.find('tokenizers/punkt')
        except LookupError:
            print("📦 Downloading NLTK punkt tokenizer...")
            nltk.download('punkt', quiet=True)
    
    async def save_uploaded_file(self, file_content: bytes, filename: str) -> str:
        """Save uploaded file and return file path"""
        file_id = str(uuid.uuid4())
        file_extension = os.path.splitext(filename)[1]
        safe_filename = f"{file_id}{file_extension}"
        file_path = os.path.join(settings.upload_dir, safe_filename)
        
        async with aiofiles.open(file_path, 'wb') as f:
            await f.write(file_content)
        
        return file_path
    
    def extract_text_from_file(self, file_path: str) -> Dict[str, Any]:
        """Advanced text extraction with structure preservation"""
        try:
            file_extension = os.path.splitext(file_path)[1].lower()
            
            if file_extension == '.pdf':
                return self.extract_from_pdf_advanced(file_path)
            elif file_extension in ['.docx', '.doc']:
                return self.extract_from_docx_advanced(file_path)
            elif file_extension in ['.txt', '.md']:
                return self.extract_from_text_advanced(file_path)
            else:
                return self.extract_from_text_advanced(file_path)  # Fallback
                
        except Exception as e:
            print(f"❌ Error extracting text: {e}")
            return {"text": "", "structured_data": {}, "metadata": {}}
    
    def extract_from_pdf_advanced(self, file_path: str) -> Dict[str, Any]:
        """Advanced PDF extraction using only pdfplumber (MIT License)"""
        extracted_data = {
            "text": "",
            "structured_data": {},
            "metadata": {"source": "pdf", "filename": os.path.basename(file_path)},
            "tables": [],
            "sections": []
        }
        
        try:
            # Use pdfplumber (MIT License - completely free)
            with pdfplumber.open(file_path) as pdf:
                full_text = ""
                tables_found = []
                
                for page_num, page in enumerate(pdf.pages):
                    # Extract text
                    page_text = page.extract_text()
                    if page_text:
                        full_text += f"\n[Page {page_num + 1}]\n{page_text}"
                    
                    # Extract tables
                    tables = page.extract_tables()
                    for table in tables:
                        if table:
                            table_text = self.table_to_text(table)
                            tables_found.append({
                                "page": page_num + 1,
                                "content": table_text,
                                "raw_data": table
                            })
                            full_text += f"\n[Table on Page {page_num + 1}]\n{table_text}\n"
                
                extracted_data["text"] = full_text
                extracted_data["tables"] = tables_found
                
        except Exception as e:
            print(f"❌ PDF extraction failed: {e}")
            # Basic fallback - try to read as text
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    extracted_data["text"] = f.read()
            except:
                pass
        
        # Extract structured data patterns
        extracted_data["structured_data"] = self.extract_structured_data(extracted_data["text"])
        
        return extracted_data
    
    def extract_from_docx_advanced(self, file_path: str) -> Dict[str, Any]:
        """Advanced DOCX extraction with structure"""
        extracted_data = {
            "text": "",
            "structured_data": {},
            "metadata": {"source": "docx", "filename": os.path.basename(file_path)},
            "tables": [],
            "sections": []
        }
        
        try:
            doc = Document(file_path)
            full_text = ""
            tables_found = []
            
            # Extract paragraphs with style information
            for para in doc.paragraphs:
                if para.text.strip():
                    style = para.style.name if para.style else "Normal"
                    if style.startswith('Heading'):
                        full_text += f"\n[{style}] {para.text}\n"
                    else:
                        full_text += f"{para.text}\n"
            
            # Extract tables
            for table_index, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                
                if table_data:
                    table_text = self.table_to_text(table_data)
                    tables_found.append({
                        "index": table_index,
                        "content": table_text,
                        "raw_data": table_data
                    })
                    full_text += f"\n[Table {table_index + 1}]\n{table_text}\n"
            
            extracted_data["text"] = full_text
            extracted_data["tables"] = tables_found
            extracted_data["structured_data"] = self.extract_structured_data(full_text)
            
        except Exception as e:
            print(f"❌ Error reading DOCX: {e}")
            
        return extracted_data
    
    def extract_from_text_advanced(self, file_path: str) -> Dict[str, Any]:
        """Advanced text file extraction"""
        extracted_data = {
            "text": "",
            "structured_data": {},
            "metadata": {"source": "text", "filename": os.path.basename(file_path)},
            "tables": [],
            "sections": []
        }
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
                extracted_data["text"] = text
                extracted_data["structured_data"] = self.extract_structured_data(text)
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    text = file.read()
                    extracted_data["text"] = text
                    extracted_data["structured_data"] = self.extract_structured_data(text)
            except Exception as e:
                print(f"❌ Error reading text file: {e}")
                
        return extracted_data
    
    def extract_structured_data(self, text: str) -> Dict[str, List[str]]:
        """Extract structured data patterns like CIN, phone numbers, emails"""
        structured_data = {}
        
        patterns = {
            'cin_numbers': r'CIN\s*:?\s*([A-Z]\d{5}[A-Z]{2}\d{4}[A-Z]{3}\d{6})',
            'phone_numbers': r'(\+?[\d\s\-\(\)]{10,15})',
            'email_addresses': r'([\w\.-]+@[\w\.-]+\.\w+)',
            'registration_numbers': r'Registration\s+(?:No\.?|Number)\s*:?\s*([A-Z0-9\-/]+)',
            'dates': r'(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})',
            'pan_numbers': r'PAN\s*:?\s*([A-Z]{5}\d{4}[A-Z])',
            'gst_numbers': r'GST\s*:?\s*(\d{2}[A-Z]{5}\d{4}[A-Z]\d[Z][A-Z\d])',
            'bank_account_numbers': r'Account\s+(?:No\.?|Number)\s*:?\s*(\d{9,18})',
            'ifsc_codes': r'IFSC\s*:?\s*([A-Z]{4}0[A-Z0-9]{6})',
        }
        
        for key, pattern in patterns.items():
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                # Clean and deduplicate matches
                cleaned_matches = list(set([match.strip() for match in matches if match.strip()]))
                structured_data[key] = cleaned_matches
        
        return structured_data
    
    def table_to_text(self, table_data: List[List[str]]) -> str:
        """Convert table data to readable text"""
        if not table_data:
            return ""
        
        text_lines = []
        for row in table_data:
            if any(cell.strip() for cell in row):  # Skip empty rows
                text_lines.append(" | ".join(cell.strip() for cell in row))
        
        return "\n".join(text_lines)
    
    def intelligent_chunking(self, text: str, metadata: Dict) -> List[Dict[str, Any]]:
        """Intelligent chunking using NLTK sentence splitting"""
        
        # Split into sentences first
        sentences = self.split_into_sentences(text)
        if len(sentences) < 3:
            return [{
                "content": text,
                "metadata": {**metadata, "chunk_method": "single_chunk"}
            }]
        
        # Group sentences into logical chunks
        chunks = []
        current_chunk = []
        current_length = 0
        
        for sentence in sentences:
            sentence_length = len(sentence)
            
            # If adding this sentence would exceed target size
            if current_length + sentence_length > self.target_chunk_size and current_chunk:
                # Save current chunk
                chunk_text = " ".join(current_chunk)
                chunks.append({
                    "content": chunk_text,
                    "metadata": {
                        **metadata,
                        "chunk_method": "intelligent_sentence",
                        "sentence_count": len(current_chunk)
                    }
                })
                
                # Start new chunk with overlap (keep last 2 sentences)
                overlap_sentences = current_chunk[-2:] if len(current_chunk) > 2 else []
                current_chunk = overlap_sentences + [sentence]
                current_length = sum(len(s) for s in current_chunk)
            else:
                current_chunk.append(sentence)
                current_length += sentence_length
        
        # Add final chunk
        if current_chunk:
            chunk_text = " ".join(current_chunk)
            chunks.append({
                "content": chunk_text,
                "metadata": {
                    **metadata,
                    "chunk_method": "intelligent_sentence",
                    "sentence_count": len(current_chunk)
                }
            })
        
        return chunks
    
    def split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences using NLTK"""
        try:
            from nltk.tokenize import sent_tokenize
            sentences = sent_tokenize(text)
            return [s.strip() for s in sentences if s.strip()]
        except Exception as e:
            print(f"⚠️ NLTK sentence splitting failed: {e}")
            # Fallback regex-based sentence splitting
            sentences = re.split(r'(?<=[.!?])\s+', text)
            return [s.strip() for s in sentences if s.strip()]
    
    def create_structured_data_chunks(self, structured_data: Dict, base_metadata: Dict) -> List[Dict[str, Any]]:
        """Create special chunks for structured data (CIN, phone numbers, etc.)"""
        structured_chunks = []
        
        for data_type, values in structured_data.items():
            for value in values:
                # Create descriptive content for better matching
                descriptions = {
                    'cin_numbers': f"Corporate Identification Number (CIN): {value}",
                    'phone_numbers': f"Phone Number: {value}",
                    'email_addresses': f"Email Address: {value}",
                    'registration_numbers': f"Registration Number: {value}",
                    'dates': f"Date: {value}",
                    'pan_numbers': f"PAN Number: {value}",
                    'gst_numbers': f"GST Number: {value}",
                    'bank_account_numbers': f"Bank Account Number: {value}",
                    'ifsc_codes': f"IFSC Code: {value}",
                }
                
                content = descriptions.get(data_type, f"{data_type}: {value}")
                
                structured_chunks.append({
                    "content": content,
                    "metadata": {
                        **base_metadata,
                        "chunk_type": "structured_data",
                        "data_type": data_type,
                        "exact_value": value,
                        "searchable_terms": [value, data_type.replace('_', ' ')]
                    }
                })
        
        return structured_chunks
    
    async def process_document(self, file_path: str, filename: str, company_id: str = "default") -> Dict[str, Any]:
        """Enhanced document processing with advanced extraction and intelligent chunking"""
        try:
            document_id = str(uuid.uuid4())
            
            print(f"📄 Processing document: {filename}")
            
            # Advanced text extraction
            extraction_result = self.extract_text_from_file(file_path)
            text_content = extraction_result["text"]
            structured_data = extraction_result["structured_data"]
            
            if not text_content.strip():
                return {"success": False, "error": "No text could be extracted from the document"}
            
            print(f"📝 Extracted {len(text_content)} characters")
            print(f"🔍 Found structured data: {list(structured_data.keys())}")
            
            # Prepare base metadata
            base_metadata = {
                **extraction_result["metadata"],
                "filename": filename,
                "company_id": company_id,
                "document_id": document_id
            }
            
            # Intelligent chunking for main content
            print("🧠 Performing intelligent chunking...")
            text_chunks = self.intelligent_chunking(text_content, base_metadata)
            
            # Create special chunks for structured data
            structured_chunks = self.create_structured_data_chunks(structured_data, base_metadata)
            
            # Combine all chunks
            all_chunks = text_chunks + structured_chunks
            
            print(f"📊 Created {len(text_chunks)} intelligent chunks + {len(structured_chunks)} structured data chunks")
            
            # Generate embeddings
            print("🔢 Generating embeddings...")
            chunk_texts = [chunk["content"] for chunk in all_chunks]
            embeddings = await self.ai_service.get_embeddings_batch(chunk_texts)
            
            # Prepare chunks for vector store
            vector_chunks = []
            for i, (chunk, embedding) in enumerate(zip(all_chunks, embeddings)):
                chunk_data = {
                    "document_id": document_id,
                    "content": chunk["content"],
                    "embedding": embedding,
                    "metadata": {
                        **chunk["metadata"],
                        "chunk_index": i,
                        "total_chunks": len(all_chunks)
                    },
                    "chunk_index": i,
                    "source_type": "custom"
                }
                vector_chunks.append(chunk_data)
            
            # Store in vector database
            print("💾 Storing in vector database...")
            success = await self.vector_store.add_documents(vector_chunks)
            
            if success:
                # Clean up uploaded file
                try:
                    os.remove(file_path)
                except:
                    pass
                
                return {
                    "success": True,
                    "document_id": document_id,
                    "chunks_processed": len(all_chunks),
                    "structured_data_found": structured_data,
                    "filename": filename,
                    "processing_method": "intelligent_chunking_nltk"
                }
            else:
                return {"success": False, "error": "Failed to store in vector database"}
                
        except Exception as e:
            print(f"❌ Error processing document: {e}")
            return {"success": False, "error": str(e)}